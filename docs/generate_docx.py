"""
============================================================
INF2003 Group 11 — Final Report DOCX Generator
Converts G11_Final_Report.md to a professionally formatted
Word document with cover page, styled headings, tables, etc.
============================================================
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Configuration ──────────────────────────────────────────
import os
DOCS_DIR = Path(os.path.dirname(os.path.abspath(__file__))) if '__file__' in dir() else Path(os.getcwd()) / "docs"
OUTPUT_PATH = DOCS_DIR / "G11_Final_Report.docx"
# Fallback if file is locked (e.g., open in Word)
try:
    OUTPUT_PATH.touch(exist_ok=True)
except PermissionError:
    import time
    OUTPUT_PATH = DOCS_DIR / f"G11_Final_Report_{int(time.time())}.docx"

# Brand colors
PRIMARY = RGBColor(0x19, 0x76, 0xD2)    # Blue (headings)
SECONDARY = RGBColor(0x0D, 0x47, 0xA1)  # Dark blue (borders)
ACCENT = RGBColor(0x2E, 0x7D, 0x32)     # Green (for special highlights)
LIGHT_BG = RGBColor(0xE3, 0xF2, 0xFD)   # Light blue (table headers)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MED_GRAY = RGBColor(0x75, 0x75, 0x75)

# ── DOCX Helpers ───────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set border of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" '
            f'w:color="{val.get("color", "1976D2")}" '
            f'w:space="0"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_paragraph(doc, text, style='Normal', bold=False, italic=False,
                         size=None, color=None, alignment=None, space_after=None,
                         space_before=None):
    """Add a paragraph with custom formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p

def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = PRIMARY
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(12)
    h.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_bullet(doc, text, level=0, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = DARK_GRAY
        run_t = p.add_run(text)
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = DARK_GRAY
    else:
        p.clear()
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.color.rgb = DARK_GRAY
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "1976D2")

    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GRAY
            if r % 2 == 1:
                set_cell_shading(cell, "F5F5F5")

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacer
    return table


def add_code_block(doc, code_text, language=""):
    """Add a monospace code block with light grey background."""
    # Label
    if language:
        label = doc.add_paragraph()
        run = label.add_run(f"▌ {language}")
        run.font.size = Pt(8)
        run.font.color.rgb = MED_GRAY
        run.bold = True
        label.paragraph_format.space_after = Pt(1)

    # Code
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        # Light grey background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
        p._element.get_or_add_pPr().append(shading)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph()  # spacer


# ── Document Builder ────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page Setup ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # ── COVER PAGE ──────────────────────────────────────
    # Add spacing to center content vertically
    for _ in range(6):
        doc.add_paragraph()

    # SIT Logo / Module line
    add_styled_paragraph(doc, "SINGAPORE INSTITUTE OF TECHNOLOGY",
                         size=11, color=MED_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=4)
    add_styled_paragraph(doc, "INF2003 Database Design & Implementation",
                         size=11, color=MED_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         space_after=24)

    # Project Title
    add_styled_paragraph(doc, "E-Commerce Clickstream &",
                         size=26, bold=True, color=PRIMARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_styled_paragraph(doc, "Transaction Analytics",
                         size=26, bold=True, color=PRIMARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    # Subtitle
    add_styled_paragraph(doc, "Dual-Database Analytics Platform",
                         size=14, italic=True, color=SECONDARY,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Horizontal rule
    p_hr = doc.add_paragraph()
    p_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:color="1976D2"/></w:pBdr>')
    pPr.append(pBdr)

    # Cover info table
    cover_data = [
        ("Group ID", "G11 — Team Hanzalians"),
        ("Topic", "Topic 2 — Dual-Database Analytics Platform"),
        ("Team Lead", "Hanzalah Hisam"),
        ("Team Members", "Faris Zharfan, Lucas Leow, Muhammad Hasan Bin Suwandi,\nMuhammad Raees Irfan Bin Ishak, Muhammad 'Afif Bin Muhd Lotfi Jarhom"),
        ("Submission Date", "13 July 2026"),
    ]

    table = doc.add_table(rows=len(cover_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(cover_data):
        cell_l = table.rows[i].cells[0]
        cell_l.width = Cm(4.5)
        run_l = cell_l.paragraphs[0].add_run(label)
        run_l.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = PRIMARY
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        cell_r = table.rows[i].cells[1]
        cell_r.width = Cm(9.0)
        run_r = cell_r.paragraphs[0].add_run(value)
        run_r.font.size = Pt(11)
        run_r.font.color.rgb = DARK_GRAY

        # Remove borders
        for cell in [cell_l, cell_r]:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            for edge in ['top', 'left', 'bottom', 'right']:
                element = parse_xml(
                    f'<w:{edge} {nsdecls("w")} w:val="nil"/>')
                tcBorders.append(element)
            tcPr.append(tcBorders)

    doc.add_paragraph()
    doc.add_paragraph()

    # Page break after cover
    doc.add_page_break()

    # ── TABLE OF CONTENTS (placeholder) ─────────────────
    add_heading_styled(doc, "Table of Contents", 1)
    toc_items = [
        "1. Introduction",
        "2. System-Level Database Integration",
        "3. Relational Database (PostgreSQL)",
        "4. NoSQL Database (MongoDB)",
        "5. Application Implementation",
        "6. Performance Evaluation",
        "7. Testing & Quality Assurance",
        "8. Constraints & Limitations",
        "9. Discussion & Reflections",
        "Appendix — Source Code & Documentation Files",
    ]
    for item in toc_items:
        add_styled_paragraph(doc, item, size=11, color=DARK_GRAY, space_after=2)

    doc.add_page_break()

    # ── SECTION 1: Introduction ─────────────────────────
    add_heading_styled(doc, "1. Introduction", 1)

    add_styled_paragraph(doc,
        "This project demonstrates a dual-database e-commerce analytics platform: "
        "PostgreSQL handles ACID transactions (orders, inventory, customers) while "
        "MongoDB handles BASE clickstream analytics (user sessions, funnel conversion, "
        "fraud detection). The Outbox Pattern provides CDC (Change Data Capture) "
        "synchronization between the two systems.", size=10)

    add_heading_styled(doc, "Key Achievements", 2)
    achievements = [
        "8 PostgreSQL tables with 4 database triggers (stock check, inventory deduction, outbox CDC, audit logging)",
        "4 MongoDB collections implementing established NoSQL patterns (Bucket, Computed, CDC Target, Cached)",
        "18 REST API endpoints with JWT authentication and role-based access control",
        "React frontend with interactive analytics dashboard (RFM pie charts, funnel bars, market basket tables)",
        "161-test automated test suite covering all endpoints, triggers, and error cases (100% pass rate)",
        "Docker Compose one-command startup with automatic data loading and database reset capabilities",
        "Cross-database fraud detection pipeline (MongoDB velocity check → PostgreSQL alert insertion)",
        "Performance benchmark suite comparing PostgreSQL vs MongoDB on 4 different workload types",
    ]
    for a in achievements:
        add_bullet(doc, a)

    add_styled_paragraph(doc,
        "Dataset: ~275,000 rows across 6 CSV files sourced from Kaggle e-commerce datasets, "
        "loaded via an automated data ingestion pipeline.", size=10, bold=False, italic=True)

    # ── SECTION 2: System-Level Database Integration ─────
    add_heading_styled(doc, "2. System-Level Database Integration", 1)

    add_heading_styled(doc, "2.1 Architecture", 2)
    add_styled_paragraph(doc,
        "FastAPI connects to PostgreSQL via SQLAlchemy (ORM + raw SQL) and to MongoDB "
        "via Motor (async driver). The Outbox Pattern bridges the two: after each order "
        "commit, a PostgreSQL trigger writes to the outbox table, and a background async "
        "poller syncs denormalized summaries to MongoDB's customer_order_summary collection.",
        size=10)

    add_heading_styled(doc, "2.2 Consistency Strategy", 2)
    add_bullet(doc, "Write Path: PostgreSQL-first (ACID) → Outbox trigger → Async poller → MongoDB (eventual consistency)", bold_prefix="Write Path: ")
    add_bullet(doc, "Read Path: Queries target the appropriate database directly — no cross-DB JOINs are performed.", bold_prefix="Read Path: ")
    add_bullet(doc, "Failure Handling: Unprocessed outbox events remain marked processed = false and are retried on the next poll cycle.", bold_prefix="Failure Handling: ")

    # ── SECTION 3: Relational Database ──────────────────
    add_heading_styled(doc, "3. Relational Database (PostgreSQL)", 1)

    # ── ER Diagram Images (Two Parts) ──────────────────
    from pathlib import Path as P
    pg_path = P(DOCS_DIR) / "ER_Diagram_PostgreSQL.png"
    mongo_path = P(DOCS_DIR) / "ER_Diagram_MongoDB.png"

    if pg_path.exists() and mongo_path.exists():
        add_heading_styled(doc, "3.0 Entity-Relationship Diagram", 2)

        # Part 1: PostgreSQL
        add_styled_paragraph(doc, "Part 1: PostgreSQL — Relational Database (8 Tables, 4 Triggers)",
                             size=11, bold=True, color=PRIMARY, space_after=4)
        img1 = doc.add_paragraph()
        img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = img1.add_run()
        run1.add_picture(str(pg_path), width=Inches(6.3))

        add_styled_paragraph(doc,
            "Figure 1a: PostgreSQL ER Diagram — 8 tables with FK relationships (solid blue arrows), "
            "4 database triggers (dashed orange arrows), and fraud detection link (dotted gray). "
            "The trigger info box details all 4 automatic database rules: stock check, inventory "
            "deduction, outbox CDC event generation, and audit logging.",
            size=8, italic=True, color=MED_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()

        # Part 2: MongoDB
        add_styled_paragraph(doc, "Part 2: MongoDB — NoSQL Document Store (4 Collections, 4 Patterns)",
                             size=11, bold=True, color=SECONDARY, space_after=4)
        img2 = doc.add_paragraph()
        img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = img2.add_run()
        run2.add_picture(str(mongo_path), width=Inches(6.3))

        add_styled_paragraph(doc,
            "Figure 1b: MongoDB Collections — 4 collections implementing established NoSQL design "
            "patterns: Bucket Pattern (user_sessions), Computed Pattern (session_stats), CDC Target "
            "Pattern (customer_order_summary), and Cached Pattern (funnel_metrics). The CDC sync box "
            "explains the Outbox Pattern flow from PostgreSQL to MongoDB. The operations box lists "
            "all MongoDB operations and indexes used.",
            size=8, italic=True, color=MED_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()

        # Explanation
        add_styled_paragraph(doc,
            "Diagram Explanation: The ER diagrams above illustrate the complete dual-database "
            "architecture split across two pages for readability. In Figure 1a (PostgreSQL), "
            "8 tables model the transactional core: users stores authentication credentials with "
            "bcrypt-hashed passwords, customers holds customer profiles with UUID primary keys, "
            "products maintains the catalog with CHECK constraints on stock quantities, orders "
            "records each purchase linked to a customer, and order_items captures individual line "
            "items with foreign keys to both orders and products. The outbox table serves as a CDC "
            "event store, order_audit_log tracks every field-level change, and alerts records fraud "
            "detection results.", size=10)

        add_styled_paragraph(doc,
            "In Figure 1b (MongoDB), 4 collections implement established NoSQL patterns: "
            "user_sessions uses the Bucket Pattern to accumulate clickstream events via atomic "
            "$push + $inc operations, session_stats provides pre-computed session aggregates "
            "(Computed Pattern), customer_order_summary holds denormalized order data synced via "
            "CDC (CDC Target Pattern), and funnel_metrics caches conversion funnel results (Cached "
            "Pattern). The CDC sync box details the 4-step Outbox Pattern flow from PostgreSQL "
            "trigger to MongoDB atomic update every 5 seconds.", size=10)
        doc.add_paragraph()

    add_heading_styled(doc, "3.1 Schema (8 Tables)", 2)
    add_styled_paragraph(doc,
        "The PostgreSQL schema consists of 8 tables with 3 relationship types and "
        "CHECK constraints enforcing data integrity at the database level:", size=10)

    pg_tables = [
        ["users", "Authentication accounts", "INT PK, username (UNIQUE), email (UNIQUE), bcrypt password_hash, role (customer/admin)"],
        ["customers", "Customer profiles", "UUID PK, country_code, opt_in_status. UUID prevents enumeration attacks."],
        ["products", "Product catalog", "VARCHAR PK, name, 7 categories, unit_price, stock_quantity (CHECK >= 0). 1,197 items."],
        ["orders", "Purchase orders", "UUID PK, FK → customers, order_date, total_amount, status. ~33,580 rows (full) / 3,000 (demo)."],
        ["order_items", "Line items", "SERIAL PK, FK → orders + products, quantity (CHECK > 0). ~75,000 rows. Triggers fire here."],
        ["outbox", "CDC event store", "SERIAL PK, aggregate_id, event_type, JSONB payload, processed flag. Grows with orders."],
        ["order_audit_log", "Change history", "SERIAL PK, FK → orders, changed_by, field_name, old_value, new_value. Trigger-populated."],
        ["alerts", "Fraud detection", "SERIAL PK, customer_id, message, alert_type, acknowledged flag. Cross-DB pipeline target."],
    ]
    add_styled_table(doc, ["Table", "Purpose", "Key Features"], pg_tables, [3, 3.5, 7])

    add_styled_paragraph(doc,
        "Relationships: customers 1:M orders → orders 1:M order_items ← products M:1 order_items. "
        "UUID primary keys on customers and orders prevent enumeration attacks. CHECK constraints "
        "enforce stock_quantity >= 0 and quantity > 0 at the database level.", size=10, italic=True)

    add_heading_styled(doc, "3.2 Triggers (4)", 2)
    triggers_data = [
        ["trg_check_stock", "order_items (BEFORE INSERT)", "Raises exception if stock_quantity < requested quantity"],
        ["trg_deduct_inventory", "order_items (AFTER INSERT)", "Decrements products.stock_quantity by order_items.quantity"],
        ["trg_outbox_order", "orders (AFTER INSERT)", "Writes JSON event to outbox table for CDC sync to MongoDB"],
        ["trg_audit_order", "orders (AFTER UPDATE)", "Logs old → new values when status or total_amount changes"],
    ]
    add_styled_table(doc, ["Trigger", "Fires On", "Purpose"], triggers_data, [3.5, 4, 6])

    add_heading_styled(doc, "3.3 Advanced SQL Queries", 2)
    add_styled_paragraph(doc,
        "The following queries demonstrate advanced relational database techniques: "
        "Common Table Expressions (CTEs), NTILE window functions, self-joins, and "
        "CASE expressions — all executed as raw SQL via SQLAlchemy's text() function.", size=10)

    add_styled_paragraph(doc,
        "RFM (Recency, Frequency, Monetary) segmentation classifies every customer into "
        "one of 5 behavioural segments. This query uses a two-stage CTE: first, order_summary "
        "aggregates per-customer metrics (last order date, order count, total spend). Second, "
        "rfm_scores applies the NTILE(4) window function to divide customers into quartiles "
        "along each dimension. The final SELECT combines the scores and maps ranges to human-"
        "readable segment labels (Champions, Loyal Customers, Potential Loyalists, At Risk, Lost). "
        "NTILE is a PostgreSQL window function that cannot be easily expressed in ORM — raw SQL "
        "was chosen deliberately to demonstrate database-level expertise.", size=9, color=MED_GRAY)

    add_code_block(doc, '''WITH order_summary AS (
    SELECT
        o.customer_id,
        MAX(o.order_date) AS last_order_date,
        COUNT(o.order_id) AS frequency,
        COALESCE(SUM(o.total_amount), 0) AS monetary
    FROM orders o
    GROUP BY o.customer_id
),
rfm_scores AS (
    SELECT
        customer_id,
        EXTRACT(DAY FROM (NOW() - last_order_date)) AS recency_days,
        frequency, monetary,
        NTILE(4) OVER (ORDER BY EXTRACT(DAY FROM (NOW() - last_order_date)) DESC) AS r_score,
        NTILE(4) OVER (ORDER BY frequency ASC) AS f_score,
        NTILE(4) OVER (ORDER BY monetary ASC) AS m_score
    FROM order_summary
)
SELECT
    customer_id::TEXT, recency_days, frequency,
    ROUND(monetary::numeric, 2) AS monetary,
    r_score, f_score, m_score,
    (r_score + f_score + m_score) AS total_score,
    CASE
        WHEN (r_score + f_score + m_score) >= 10 THEN 'Champions'
        WHEN (r_score + f_score + m_score) >= 8  THEN 'Loyal Customers'
        WHEN (r_score + f_score + m_score) >= 6  THEN 'Potential Loyalists'
        WHEN (r_score + f_score + m_score) >= 4  THEN 'At Risk'
        ELSE 'Lost'
    END AS segment
FROM rfm_scores
ORDER BY total_score DESC;''', "RFM Segmentation — CTE + NTILE(4) + CASE (relational_service.py)")

    add_code_block(doc, '''SELECT
    p1.product_id AS product_a, pa.name AS name_a,
    p2.product_id AS product_b, pb.name AS name_b,
    COUNT(*) AS pair_count
FROM order_items p1
JOIN order_items p2
    ON p1.order_id = p2.order_id
    AND p1.product_id < p2.product_id
JOIN products pa ON pa.product_id = p1.product_id
JOIN products pb ON pb.product_id = p2.product_id
GROUP BY p1.product_id, pa.name, p2.product_id, pb.name
ORDER BY pair_count DESC
LIMIT :top_n;''', "Market Basket Analysis — Self-Join on order_items (relational_service.py)")

    add_styled_paragraph(doc,
        "Market Basket Analysis discovers which products are frequently purchased together — "
        "a technique used by Amazon ('Customers who bought this also bought...'). This query "
        "performs a self-join on the order_items table: each row is paired with every other row "
        "from the same order. The condition p1.product_id < p2.product_id prevents duplicate "
        "pairs (A+B is the same as B+A). Two additional JOINs to the products table resolve "
        "product names for readability. Results are sorted by pair frequency — the top pair "
        "represents the strongest product affinity.", size=9, color=MED_GRAY)

    # ── SECTION 4: NoSQL Database ───────────────────────
    add_heading_styled(doc, "4. NoSQL Database (MongoDB)", 1)

    add_heading_styled(doc, "4.1 Collections (4)", 2)
    mongo_data = [
        ["user_sessions", "Bucket", "Accumulates clickstream events per session via $push + $inc. ~27,500 documents."],
        ["session_stats", "Computed", "Pre-aggregated session summaries for fast dashboard queries."],
        ["customer_order_summary", "CDC Target", "Denormalized view synced from PostgreSQL via Outbox. Updated with $inc."],
        ["funnel_metrics", "Cached", "Aggregated conversion funnel results computed via $facet pipeline."],
    ]
    add_styled_table(doc, ["Collection", "Pattern", "Purpose"], mongo_data, [3.5, 2.5, 7.5])

    add_heading_styled(doc, "4.2 Advanced NoSQL Queries", 2)
    add_styled_paragraph(doc,
        "MongoDB's aggregation pipeline demonstrates document-oriented analytics "
        "without relational JOINs. The $facet stage runs 4 parallel counts in a "
        "single pass through the data.", size=10)

    add_styled_paragraph(doc,
        "This $facet aggregation pipeline computes the e-commerce conversion funnel in a single "
        "database pass. First, $unwind expands each session's embedded events[] array into "
        "individual documents. Then $facet runs four independent sub-pipelines in parallel: "
        "each filters for a specific action_type (page_view, add_to_cart, checkout, purchase), "
        "groups by session_id to deduplicate sessions, and counts unique sessions per stage. "
        "The backend then calculates conversion rates by dividing each stage count by the "
        "page_view total. This replaces what would require 4 separate SQL queries with "
        "a single MongoDB operation.", size=9, color=MED_GRAY)

    add_code_block(doc, '''db.user_sessions.aggregate([
  { $unwind: "$events" },
  { $facet: {
      page_views: [
        { $match: {"events.action_type": "page_view"}},
        { $group: {_id: "$session_id"}},
        { $count: "count"}
      ],
      add_to_cart: [
        { $match: {"events.action_type": "add_to_cart"}},
        { $group: {_id: "$session_id"}},
        { $count: "count"}
      ],
      checkouts: [
        { $match: {"events.action_type": "checkout"}},
        { $group: {_id: "$session_id"}},
        { $count: "count"}
      ],
      purchases: [
        { $match: {"events.action_type": "purchase"}},
        { $group: {_id: "$session_id"}},
        { $count: "count"}
      ]
  }}
])''', "Funnel Analytics — MongoDB $facet Aggregation Pipeline (nosql_service.py)")

    add_code_block(doc, '''# Velocity-based fraud detection (nosql_service.py)
threshold = settings.FRAUD_EVENT_THRESHOLD  # default: 10
window_start = datetime.utcnow() - timedelta(
    seconds=settings.FRAUD_TIME_WINDOW_SECONDS)

session = await db.user_sessions.find_one({
    "customer_id": customer_id, "session_id": session_id
})

recent_events = [e for e in session.get("events", [])
                 if e["timestamp"] >= window_start]

if len(recent_events) >= threshold:
    has_purchase = any(e["action_type"] == "purchase"
                       for e in recent_events)
    if not has_purchase:
        await db.user_sessions.update_one(
            {"customer_id": customer_id, "session_id": session_id},
            {"$set": {"flagged": True}})
        # Cross-database alert: MongoDB flag -> PostgreSQL INSERT
        create_alert(db, customer_id, session_id, ...)''', "Fraud Detection — Cross-Database Velocity Check (nosql_service.py)")

    add_styled_paragraph(doc,
        "This velocity-based fraud detector fires whenever a user exceeds the configured threshold "
        "(default: 10 add_to_cart events in 60 seconds) without completing a purchase. First, it "
        "retrieves the user's session document from MongoDB (O(1) lookup by compound index on "
        "customer_id + session_id). It then filters events to the current time window and checks "
        "whether any purchase event exists. If the threshold is exceeded with no purchase, it "
        "flags the session in MongoDB ($set: {flagged: True}) and INSERTs an alert into "
        "PostgreSQL — demonstrating cross-database coordination where the high-velocity event "
        "store (MongoDB) feeds the authoritative alert log (PostgreSQL).", size=9, color=MED_GRAY)

    # ── SECTION 5: Application Implementation ───────────
    add_heading_styled(doc, "5. Application Implementation", 1)

    add_heading_styled(doc, "5.1 Web Interface (React + Recharts)", 2)
    add_styled_paragraph(doc,
        "The frontend is a single-page application built with React 18 and Vite, "
        "featuring 4 functional views:", size=10)

    web_data = [
        ["Product Catalog", "/", "1,197 products across 7 categories, search, filter, pagination (20/page, 60 pages)"],
        ["Cart & Clickstream", "/cart", "4 event types, real-time event table, session tracking, fraud trigger demo"],
        ["Admin Dashboard", "/admin", "RFM pie chart, funnel bar chart, market basket, top products, alerts (admin only)"],
        ["Login/Register", "/login", "JWT auth with bcrypt hashing, form validation, role-based redirect"],
    ]
    add_styled_table(doc, ["Page", "Route", "Features"], web_data, [3, 2, 8.5])

    add_heading_styled(doc, "5.2 Key Features", 2)
    features = [
        "JWT authentication with bcrypt password hashing and role-based access control (customer/admin). The first user (user_1) is auto-assigned admin role for instant dashboard access.",
        "Product catalog with category filter (7 categories), text search, product names, and pagination (20/100 per page)",
        "Clickstream event tracking via MongoDB Bucket Pattern — all 4 funnel event types recorded. Bulk data loading uses batch $push: { $each: [...] } (one DB call per session); runtime uses individual $push + $inc (O(1) per event)",
        "Order creation with full ACID transaction — FK validation, stock check, inventory deduction, outbox CDC, and audit logging all fire automatically via database triggers",
        "Real-time fraud detection — MongoDB session velocity analysis triggers PostgreSQL alert insertion when 10+ add_to_cart events occur in 60 seconds without a purchase",
        "Admin dashboard with RFM pie chart, funnel bar chart (4-stage conversion), market basket (top 10 pairs with product names), top products, sales by category, and fraud alerts",
        "Optimized data loader: ~50s demo mode (24x speedup via batch MongoDB writes, nrows sampling, shared bcrypt hash). Controlled via DEMO_MODE env var in .env file. Full 275K-row dataset via DEMO_MODE=false",
        "Automated test suite — 161 tests covering every endpoint, trigger, and error case (3-7 second runtime, 100% pass)",
        "Docker Compose one-command startup with automatic data loading and database reset capabilities",
    ]
    for f in features:
        add_bullet(doc, f)

    # ── SECTION 6: Performance Evaluation ───────────────
    add_heading_styled(doc, "6. Performance Evaluation", 1)
    add_styled_paragraph(doc,
        "Benchmarks were run via backend/benchmark/benchmark_runner.py inside Docker "
        "on 2026-06-30. The suite measures 4 distinct workload types to highlight "
        "the strengths of each database:", size=10)

    bench_data = [
        ["1", "Bulk Insert (10k events)", "MongoDB", "14.39s", "695 ops/sec", "Bucket Pattern avoids per-event overhead"],
        ["2", "Hotspot UPDATEs (200 txns)", "PostgreSQL", "0.53s", "378 txns/sec", "Row-level locking under contention"],
        ["3", "5-Table JOIN", "PostgreSQL", "0.019s", "—", "Query optimizer handles complex algebra"],
        ["4", "Aggregation Pipeline", "MongoDB", "0.117s", "—", "$facet: 4 parallel counts in single pass"],
    ]
    add_styled_table(doc, ["#", "Test", "Database", "Time", "Throughput", "Insight"],
                     bench_data, [0.8, 3.2, 1.8, 1.5, 1.8, 4.5])

    add_styled_paragraph(doc,
        "Key finding: PostgreSQL's 5-table JOIN (0.019s) outperforms MongoDB's "
        "aggregation pipeline (0.117s) for complex relational queries — despite "
        "MongoDB not needing JOINs at all. This validates the polyglot persistence "
        "approach: PostgreSQL for structured queries with relationships, MongoDB "
        "for high-velocity clickstream writes (695 ops/sec via Bucket Pattern).", size=10)

    # Insert benchmark plot image
    plot_path = DOCS_DIR / ".." / "backend" / "benchmark" / "plots" / "benchmark_results.png"
    if plot_path.exists():
        add_styled_paragraph(doc, "Benchmark Results Chart:", size=10, bold=True, space_after=4)
        try:
            img = doc.add_picture(str(plot_path), width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            add_styled_paragraph(doc, "(Chart image could not be embedded — see benchmark/plots/benchmark_results.png)", size=9, color=MED_GRAY)
    else:
        add_styled_paragraph(doc, "(Benchmark plot not found — run benchmark_runner.py first)", size=9, color=MED_GRAY)

    # ── SECTION 7: Testing & QA ─────────────────────────
    add_heading_styled(doc, "7. Testing & Quality Assurance", 1)
    add_styled_paragraph(doc,
        "A comprehensive test suite (backend/tests/test_suite.py) was developed with "
        "161 automated tests. Result: 161/161 passed (100%) in 7.08 seconds.", size=10)

    test_data = [
        ["Health Checks", "7", "Root, PostgreSQL/MongoDB connectivity, Swagger UI, OpenAPI schema"],
        ["Authentication", "12", "Registration, login, /me, duplicates, invalid passwords, expired tokens"],
        ["Products API", "30", "Listing, pagination, category filter, search, get-by-ID, 404, categories"],
        ["Cart/Clickstream", "16", "All 4 event types, invalid inputs, session retrieval, auto-generated IDs"],
        ["Orders (ACID)", "18", "Order creation, retrieval, insufficient stock, inventory deduction, outbox CDC"],
        ["Analytics", "24", "RFM (5 segments), market basket, funnel ($facet), cart abandonment, top products"],
        ["Admin Analytics", "7", "Alerts (admin + non-admin 403), audit trail (admin + non-admin + unauthenticated)"],
        ["Trigger Verification", "10", "All 4 trigger functions confirmed, table attachments, CHECK constraints"],
        ["MongoDB Verification", "10", "All 4 collections, compound index, TTL index, flagged, timestamp indexes"],
        ["Error Handling", "6", "Empty body, short username, invalid email, short password, 404, bad pagination"],
    ]
    add_styled_table(doc, ["Section", "Tests", "Key Checks"], test_data, [3, 1.5, 9])

    # ── SECTION 8: Constraints ──────────────────────────
    add_heading_styled(doc, "8. Constraints & Limitations", 1)
    constraints = [
        ("Outbox polling introduces ~5 second staleness ", "for CDC data in MongoDB. Real-time CDC would require Kafka + Debezium, which was beyond the project scope."),
        ("Fraud detection uses simple threshold counting ", "(10 events in 60 seconds). Production systems would use ML models trained on historical fraud patterns."),
        ("No distributed tracing or production observability ", "(no OpenTelemetry, no centralized logging)."),
        ("Single-instance databases ", "— no replication, sharding, or failover configured."),
        ("The users and customers tables have no FK relationship, ", "meaning any authenticated user can place orders under any customer ID. A production system would link them."),
    ]
    for bold_part, normal_part in constraints:
        add_bullet(doc, normal_part, bold_prefix=bold_part)

    # ── SECTION 9: Discussion & Reflections ─────────────
    add_heading_styled(doc, "9. Discussion & Reflections", 1)

    add_heading_styled(doc, "9.1 What Worked Well", 2)
    worked_well = [
        ("Clear separation of concerns: ", "PostgreSQL handles everything requiring ACID (orders, inventory); MongoDB handles everything requiring high write throughput (clickstream, sessions, funnel). Neither database is a bottleneck for the other's workload."),
        ("Outbox Pattern provided reliable CDC ", "without external infrastructure (Kafka, Debezium). The pattern was simple to implement, easy to debug, and guaranteed at-least-once delivery."),
        ("Database triggers simplified the application layer ", "— inventory management, CDC event generation, and audit logging happen automatically at the database level, reducing application code complexity and eliminating race conditions."),
        ("MongoDB Bucket Pattern ", "perfectly matched the clickstream use case. Bulk loading uses $push: { $each: [...] } (one DB call per session); runtime $push + $inc operations are O(1) and avoid the document-per-event anti-pattern."),
        ("Docker Compose ", "made the project instantly reproducible. The data loader auto-runs on first start, and the reset service enables clean demos."),
    ]
    for bold_part, normal_part in worked_well:
        add_bullet(doc, normal_part, bold_prefix=bold_part)

    add_heading_styled(doc, "9.2 Challenges", 2)
    challenges = [
        ("Cross-database consistency ", "required careful design. The Outbox Pattern provides eventual consistency, but debugging sync failures required checking both PostgreSQL (outbox.processed) and MongoDB (customer_order_summary)."),
        ("Trigger debugging was difficult ", "— PostgreSQL trigger errors roll back the entire transaction, and error messages appear only in server logs, not application output."),
        ("The JWT sub claim initially used an integer ", "user_id, which python-jose rejects (requires string). This was discovered and fixed during integration testing (changed to str(user.user_id))."),
        ("Bcrypt password hashing was a significant bottleneck ", "— an early data loader version called pwd_ctx.hash() inside a loop, which would have taken ~83 minutes for 20,000 users. Fixed by computing the hash once outside the loop and reusing it. Additionally, a hardcoded bcrypt hash was incompatible with the container's bcrypt library version — resolved by switching to runtime hash generation via passlib."),
        ("Product names contained trailing synthetic numbers ", "from the Kaggle dataset (e.g., 'SSD MediumBlue 149'). These were cleaned using regex (re.sub(r'\\s+\\d+$', '', name)) in the data loader. A name column was added to the products table and propagated through all API responses and the frontend."),
        ("Vite HMR with Docker volumes was unreliable on Windows ", "— Vite's Hot Module Replacement did not consistently detect file changes on Docker volume mounts. Workaround: restart the frontend container (docker restart ecommerce-frontend) after code changes."),
    ]
    for bold_part, normal_part in challenges:
        add_bullet(doc, normal_part, bold_prefix=bold_part)

    add_heading_styled(doc, "9.3 GenAI Usage Reflection", 2)
    add_styled_paragraph(doc,
        "GenAI tools (ChatGPT, DeepSeek) were used throughout the project to generate initial "
        "boilerplate for FastAPI routes, SQLAlchemy models, and React components; debug PostgreSQL "
        "trigger syntax and MongoDB aggregation pipelines; and draft documentation (README, walkthrough, "
        "ER diagram).", size=10)
    add_bullet(doc, "Significantly accelerated development; provided working examples for unfamiliar technologies (Motor async driver, Recharts integration); helped identify the JWT sub type bug.", bold_prefix="Pros: ")
    add_bullet(doc, "Generated code occasionally contained subtle bugs (e.g., the integer sub issue); required careful review and testing; sometimes produced over-engineered solutions.", bold_prefix="Cons: ")
    add_bullet(doc, "Use GenAI for boilerplate generation and debugging, but maintain a strong test suite as a safety net. The 161-test suite was essential for catching AI-generated bugs.", bold_prefix="Recommendation: ")

    add_heading_styled(doc, "9.4 Future Improvements", 2)
    improvements = [
        "Add Kafka + Debezium for real-time CDC (sub-second latency vs current 5-second polling)",
        "Implement Redis caching for the product catalog (reducing PostgreSQL query load)",
        "Deploy to AWS RDS (PostgreSQL) + MongoDB Atlas for cloud-native scalability",
        "Add proper user-customer linking via FK from users to customers",
        "Implement session replay — reconstruct a user's browsing journey from MongoDB events",
        "Add OAuth2 social login (Google, GitHub) alongside current username/password auth",
    ]
    for imp in improvements:
        add_bullet(doc, imp)

    # ── SECTION 10: References ──────────────────────────
    doc.add_page_break()
    add_heading_styled(doc, "10. References", 1)

    refs = [
        ("[1] ", "Waqi (2024). \"E-commerce Clickstream and Transaction Dataset.\" Kaggle. ", ""),
        ("    ", "https://www.kaggle.com/datasets/waqi786/e-commerce-clickstream-and-transaction-dataset", ""),
        ("    ", "~200,000 clickstream events + transaction records used for the sessions.csv, clickstream_events.csv, and orders.csv base data.", ""),
        ("", "", ""),
        ("[2] ", "Wafaa Elhusseini (2024). \"Synthetic E-commerce Transactions + Clickstream 2020–2025.\" Kaggle. ", ""),
        ("    ", "https://www.kaggle.com/datasets/wafaaelhusseini/e-commerce-transactions-clickstream", ""),
        ("    ", "~75,000 synthetic transactions used for customers.csv and products.csv base data.", ""),
        ("", "", ""),
        ("[3] ", "Kleppmann, M. (2017). Designing Data-Intensive Applications. O'Reilly Media.", ""),
        ("    ", "Reference for ACID vs BASE trade-offs, Outbox Pattern, and polyglot persistence architecture.", ""),
        ("", "", ""),
        ("[4] ", "MongoDB, Inc. (2024). \"Aggregation Pipeline.\" MongoDB Documentation.", ""),
        ("    ", "https://www.mongodb.com/docs/manual/core/aggregation-pipeline/", ""),
        ("    ", "Reference for $facet, $unwind, and Bucket Pattern implementation.", ""),
        ("", "", ""),
        ("[5] ", "PostgreSQL Global Development Group. (2024). \"CREATE TRIGGER.\" PostgreSQL Documentation.", ""),
        ("    ", "https://www.postgresql.org/docs/15/sql-createtrigger.html", ""),
        ("    ", "Reference for BEFORE/AFTER INSERT triggers, CHECK constraints, and CTE syntax.", ""),
    ]
    for bold_part, normal_part, extra in refs:
        if bold_part == "":
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        if bold_part.strip():
            run_b = p.add_run(bold_part)
            run_b.bold = True
            run_b.font.size = Pt(9)
        run_n = p.add_run(normal_part)
        run_n.font.size = Pt(9)
        if extra:
            p2 = doc.add_paragraph()
            run_e = p2.add_run(extra)
            run_e.font.size = Pt(8)
            run_e.font.color.rgb = MED_GRAY
            run_e.italic = True

    # ── APPENDIX ────────────────────────────────────────
    doc.add_page_break()
    add_heading_styled(doc, "Appendix A — Complete Source Code Listing", 1)
    add_styled_paragraph(doc, "This appendix is not counted toward the 8-page limit.", size=9, italic=True, color=MED_GRAY)

    appendix_data = [
        ["backend/main.py", "FastAPI entry point: lifespan hooks, CORS, router registration"],
        ["backend/config.py", "Centralised settings from environment variables (12-Factor App)"],
        ["backend/data_loader.py", "Dual-database CSV ingestion: batch $push $each, nrows sampling, shared bcrypt"],
        ["backend/triggers.sql", "5 PostgreSQL table DDL + 4 trigger functions (stock, inventory, outbox, audit)"],
        ["backend/reset_db.py", "Database wipe & recreate: drops all PG tables + MongoDB collections"],
        ["backend/promote_admin.py", "CLI tool to promote a user to admin role"],
        ["backend/models/relational.py", "8 SQLAlchemy ORM models: User, Customer, Product, Order, OrderItem, etc."],
        ["backend/models/nosql_schemas.py", "Pydantic models: UserSession, ClickstreamEvent, ActionType enum"],
        ["backend/services/relational_service.py", "Complex SQL: RFM (CTE+NTILE), market basket (self-join), top products"],
        ["backend/services/nosql_service.py", "MongoDB: Bucket Pattern, $facet funnel, fraud detection, CDC target, indexes"],
        ["backend/services/sync_service.py", "CDC Outbox processor: async poller (5s) → PG outbox → MongoDB sync"],
        ["backend/api/auth.py", "JWT: registration, login, /me with bcrypt + OAuth2PasswordBearer"],
        ["backend/api/products.py", "Product catalog: paginated list, category filter, ILIKE search"],
        ["backend/api/cart.py", "Clickstream events: MongoDB Bucket Pattern writes + fraud detection"],
        ["backend/api/orders.py", "ACID order creation with 4-trigger cascade + order history"],
        ["backend/api/analytics.py", "9 analytics endpoints: RFM, funnel, market basket, alerts, audit"],
        ["backend/tests/test_suite.py", "161-test automated suite (10 sections, 100% pass rate)"],
        ["backend/benchmark/benchmark_runner.py", "4 database benchmarks + matplotlib chart generation"],
        ["frontend/src/App.jsx", "Root React component: routing, auth state, cart management"],
        ["frontend/src/api.js", "Axios API client: 13 endpoint functions + JWT header injection"],
        ["frontend/src/components/ProductList.jsx", "Product catalog: search, filter, pagination, add-to-cart"],
        ["frontend/src/components/Cart.jsx", "Shopping cart: quantity controls, checkout, clickstream debug"],
        ["frontend/src/components/Login.jsx", "Login/Register form with URL-based mode toggle"],
        ["frontend/src/components/AdminDashboard.jsx", "Admin analytics: RFM pie, funnel bar, market basket, alerts"],
        ["docker-compose.yml", "6-service orchestration with health checks, volumes, and profiles"],
        [".env", "Environment variables template (not committed to git)"],
    ]
    add_styled_table(doc, ["File", "Description"], appendix_data, [5.5, 8])

    # ── Appendix B: API Reference ─────────────────────
    doc.add_page_break()
    add_heading_styled(doc, "Appendix B — API Endpoint Reference", 1)

    api_data = [
        ["POST", "/api/auth/register", "None", "PG", "Create account (bcrypt hash)"],
        ["POST", "/api/auth/login", "None", "PG", "Authenticate → JWT token"],
        ["GET", "/api/auth/me", "JWT", "PG", "Current user profile"],
        ["GET", "/api/products/", "None", "PG", "Paginated catalog (category, search)"],
        ["GET", "/api/products/{id}", "None", "PG", "Single product detail"],
        ["GET", "/api/products/categories/all", "None", "PG", "Distinct category list"],
        ["POST", "/api/cart/event", "JWT", "Mongo", "Record clickstream event"],
        ["GET", "/api/cart/session/{id}", "JWT", "Mongo", "Retrieve session events"],
        ["POST", "/api/orders/", "JWT", "PG", "Create order (ACID + triggers)"],
        ["GET", "/api/orders/{id}", "JWT", "PG", "Order detail by ID"],
        ["GET", "/api/orders/", "JWT", "PG", "User's order history"],
        ["GET", "/api/analytics/rfm", "JWT", "PG", "RFM customer segmentation"],
        ["GET", "/api/analytics/market-basket", "JWT", "PG", "Product affinity pairs"],
        ["GET", "/api/analytics/funnel", "JWT", "Mongo", "Conversion funnel (4-stage)"],
        ["GET", "/api/analytics/cart-abandonment", "JWT", "Mongo", "Abandoned sessions"],
        ["GET", "/api/analytics/top-products", "JWT", "PG", "Best-selling products"],
        ["GET", "/api/analytics/sales-by-category", "JWT", "PG", "Revenue by category"],
        ["GET", "/api/analytics/alerts", "Admin", "PG", "Fraud alerts (admin only)"],
        ["GET", "/api/analytics/audit/{id}", "Admin", "PG", "Order change history (admin only)"],
    ]
    add_styled_table(doc, ["Method", "Endpoint", "Auth", "DB", "Purpose"], api_data, [1.2, 4.5, 1.2, 1.2, 5.5])

    # ── Appendix C: Environment Variables ──────────────
    add_heading_styled(doc, "Appendix C — Environment Variables", 1)

    env_data = [
        ["DATABASE_URL", "postgresql://...", "PostgreSQL connection string"],
        ["MONGO_URI", "mongodb://localhost:27017", "MongoDB connection URI"],
        ["MONGO_DB", "ecommerce_nosql", "MongoDB database name"],
        ["JWT_SECRET_KEY", "(dev default)", "HMAC-SHA256 signing key"],
        ["JWT_ALGORITHM", "HS256", "JWT signing algorithm"],
        ["JWT_EXPIRE_MINUTES", "60", "Token expiry duration"],
        ["FRAUD_EVENT_THRESHOLD", "10", "Cart events to trigger fraud alert"],
        ["FRAUD_TIME_WINDOW_SECONDS", "60", "Fraud detection time window"],
        ["OUTBOX_POLL_INTERVAL", "5", "CDC poller interval (seconds)"],
        ["DEMO_MODE", "true", "Toggle demo/full dataset loading"],
        ["DEBUG", "false", "SQLAlchemy query logging"],
    ]
    add_styled_table(doc, ["Variable", "Default", "Purpose"], env_data, [4.5, 3, 6])

    # ── Appendix D: Technology Stack ───────────────────
    add_heading_styled(doc, "Appendix D — Technology Stack", 1)

    tech_data = [
        ["Frontend", "React 18, Vite 5, Recharts 2, React Router 6, Axios 1", "SPA with 4 views, interactive charts, JWT interceptor"],
        ["Backend", "Python 3.11, FastAPI, SQLAlchemy 2, Motor 3, Pandas 2", "Async REST API, ORM + raw SQL, async MongoDB driver"],
        ["Auth", "python-jose 3, passlib 1, bcrypt", "JWT encode/decode, password hashing, OAuth2 flow"],
        ["Databases", "PostgreSQL 15, MongoDB 7", "Relational (ACID) + Document store (BASE)"],
        ["Infrastructure", "Docker 24, Docker Compose v2", "6-service orchestration, health checks, volumes"],
        ["Testing", "Python requests, unittest-style", "161 automated tests, 10 sections, 100% pass rate"],
    ]
    add_styled_table(doc, ["Layer", "Technologies", "Purpose"], tech_data, [2, 5.5, 6])

    # ── Footer ──────────────────────────────────────────
    doc.add_paragraph()
    p_footer = add_styled_paragraph(doc,
        "INF2003 Group 11 — Team Hanzalians — Singapore Institute of Technology — July 2026",
        size=9, color=MED_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

    # ── Save ────────────────────────────────────────────
    try:
        doc.save(str(OUTPUT_PATH))
        print(f"DOCX saved to: {OUTPUT_PATH}")
        print(f"File size: {OUTPUT_PATH.stat().st_size / 1024:.1f} KB")
    except PermissionError:
        import time
        alt_path = DOCS_DIR / f"G11_Final_Report_{int(time.time())}.docx"
        doc.save(str(alt_path))
        print(f"DOCX saved to: {alt_path} (original was locked)")
        print(f"File size: {alt_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build_document()
